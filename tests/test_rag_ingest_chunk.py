"""Tests for F9 (ingestion) and F10 (chunking)."""
import io

from app.engine.executor import WorkflowExecutor
from app.models.workflow import WorkflowDef
from app.rag.chunking import chunk_document
from app.rag.ingestion import detect_format, ingest
from app.rag.models import Document


# --------------------------------------------------------------------------- F9
def test_f9_detect_format():
    assert detect_format("a.md") == "md"
    assert detect_format("a.PDF") == "pdf"
    assert detect_format("a.unknown") == "text"
    assert detect_format("a.txt", explicit="html") == "html"


def test_f9_ingest_html_strips_tags_and_finds_title():
    html = b"<html><head><title>Hello</title><style>x{}</style></head>" \
           b"<body><p>First para.</p><p>Second   para.</p></body></html>"
    doc = ingest(html, filename="page.html")
    assert doc.format == "html"
    assert doc.metadata["title"] == "Hello"
    assert "First para." in doc.text
    assert "{}" not in doc.text  # style dropped
    assert "  " not in doc.text  # whitespace normalized


def test_f9_ingest_markdown_metadata():
    doc = ingest(b"# Title\n\nBody text here.", filename="d.md")
    assert doc.format == "markdown"
    assert doc.metadata["title"] == "Title"
    assert doc.metadata["word_count"] > 0
    assert doc.metadata["token_estimate"] >= 1


def test_f9_ingest_json_flattens():
    doc = ingest(b'{"name": "acme", "items": [1, 2]}', filename="d.json")
    assert "name: acme" in doc.text
    assert doc.metadata["structured"] is True


def test_f9_ingest_docx_roundtrip():
    import docx

    d = docx.Document()
    d.add_paragraph("Hello from docx.")
    d.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    d.save(buf)
    doc = ingest(buf.getvalue(), filename="file.docx")
    assert doc.format == "docx"
    assert "Hello from docx." in doc.text
    assert "Second paragraph." in doc.text


# -------------------------------------------------------------------------- F10
def _doc(text):
    return Document(document_id="d1", text=text, format="text", metadata={"title": "T"})


def test_f10_fixed_overlap_covers_text():
    doc = _doc("abcdefghij" * 5)  # 50 chars
    chunks = chunk_document(doc, "fixed", {"chunk_size": 20, "overlap": 5, "size_unit": "chars"})
    assert len(chunks) >= 3
    # Overlap: each chunk after the first starts before the previous one ended.
    assert chunks[1].start_char < chunks[0].end_char
    # Provenance carried through.
    assert chunks[0].metadata["title"] == "T"
    assert chunks[0].chunk_id == "d1::0"


def test_f10_recursive_respects_size():
    text = "\n\n".join(f"Paragraph number {i} with some words." for i in range(10))
    chunks = chunk_document(_doc(text), "recursive", {"chunk_size": 60, "size_unit": "chars"})
    assert len(chunks) > 1
    assert all(c.char_count <= 60 for c in chunks)


def test_f10_semantic_keeps_sentences_whole():
    text = "First sentence here. Second sentence follows. Third one ends it."
    chunks = chunk_document(_doc(text), "semantic", {"chunk_size": 40, "size_unit": "chars"})
    # No chunk should cut a sentence in the middle (each ends with terminal punct).
    assert all(c.text.strip().endswith((".", "!", "?")) for c in chunks)


def test_f10_structure_attaches_headings():
    md = "# Intro\nHello world.\n\n## Details\nMore content here for the section."
    chunks = chunk_document(_doc(md), "structure", {"chunk_size": 200, "size_unit": "chars"})
    headings = {c.metadata.get("heading") for c in chunks}
    assert "Intro" in headings and "Details" in headings


# ------------------------------------------------------------- end-to-end node
async def test_f9_f10_pipeline_node_flow():
    wf = WorkflowDef(
        name="pipe",
        nodes=[
            {"id": "in", "type": "input"},
            {"id": "ingest", "type": "ingest",
             "config": {"text": "$.inputs.body", "filename": "g.md"}},
            {"id": "chunk", "type": "chunk",
             "config": {"strategy": "structure", "chunk_size": 80, "size_unit": "chars"}},
            {"id": "out", "type": "output", "config": {"value": "$.chunk"}},
        ],
        edges=[
            {"source": "in", "target": "ingest"},
            {"source": "ingest", "target": "chunk"},
            {"source": "chunk", "target": "out"},
        ],
    )
    body = "# A\nAlpha section text.\n\n# B\nBravo section text."
    res = await WorkflowExecutor(wf).run({"body": body})
    assert res.status == "success"
    chunks = res.outputs["out"]
    assert len(chunks) >= 2
    assert chunks[0]["document_id"] == chunks[-1]["document_id"]
    assert {c["metadata"].get("heading") for c in chunks} >= {"A", "B"}
